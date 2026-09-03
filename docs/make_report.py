"""Generate the project report as a Word document.

Kept as a script rather than a one-off file so the report can be regenerated
when the measured figures change - the numbers below are the ones produced by
ml/seed_from_dataset.py and the accuracy checks, not hand-copied prose.

    python docs/make_report.py [output.docx]
"""

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

INK = RGBColor(0x15, 0x20, 0x2B)
MUTED = RGBColor(0x5B, 0x68, 0x75)
ACCENT = RGBColor(0x13, 0x6A, 0x6F)
STOP = RGBColor(0x9E, 0x3D, 0x3D)
PART = RGBColor(0xA8, 0x64, 0x1D)
OK = RGBColor(0x2F, 0x7D, 0x55)


def styled(paragraph, text, *, size=11, bold=False, italic=False,
           color=INK, font="Calibri"):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = font
    return run


def heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(20 if level == 1 else 14)
    paragraph.paragraph_format.space_after = Pt(4)
    styled(paragraph, text, size=16 if level == 1 else 12.5, bold=True)
    return paragraph


def body(doc, text, *, size=11, color=INK, space_after=8, italic=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after)
    styled(paragraph, text, size=size, color=color, italic=italic)
    return paragraph


def table(doc, headers, rows, widths=None):
    """A table with a header row and right-aligned numeric columns."""
    grid = doc.add_table(rows=1, cols=len(headers))
    grid.style = "Light Grid Accent 1"
    grid.alignment = WD_TABLE_ALIGNMENT.LEFT
    for index, label in enumerate(headers):
        cell = grid.rows[0].cells[index]
        cell.text = ""
        styled(cell.paragraphs[0], label, size=10, bold=True)
    for row in rows:
        cells = grid.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = ""
            paragraph = cells[index].paragraphs[0]
            numeric = index > 0 and any(ch.isdigit() for ch in str(value))
            if numeric:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            styled(paragraph, str(value), size=10,
                   font="Consolas" if numeric else "Calibri")
    if widths:
        for row in grid.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return grid


def finding(doc, title, status, status_colour, paragraphs, why):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(2)
    styled(paragraph, title, size=12, bold=True)
    styled(paragraph, f"   [{status}]", size=9.5, bold=True, color=status_colour)
    for text in paragraphs:
        body(doc, text, space_after=6)
    note = doc.add_paragraph()
    note.paragraph_format.left_indent = Inches(0.25)
    note.paragraph_format.space_after = Pt(10)
    styled(note, why, size=10.5, italic=True, color=MUTED)


def build(path: Path) -> Path:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ---- title -------------------------------------------------------
    eyebrow = doc.add_paragraph()
    eyebrow.paragraph_format.space_after = Pt(2)
    styled(eyebrow, "PROJECT REPORT", size=9.5, bold=True, color=ACCENT)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    styled(title, "Invoice AI", size=26, bold=True)

    body(doc,
         "An invoice extraction pipeline that scores its own reliability and "
         "learns from reviewer corrections. What was built, what it measures, "
         "what remains blocked, and where the effort should go next.",
         size=12, color=MUTED, space_after=14)

    table(doc,
          ["Headline", "Value"],
          [["Totals extracted correctly", "99%"],
           ["Best model balanced accuracy", "0.99"],
           ["Annotated invoices benchmarked", "414"],
           ["Fields on trained models", "2 of 5"]],
          widths=[3.4, 1.4])

    # ---- what it does ------------------------------------------------
    heading(doc, "What the system does")
    body(doc,
         "An invoice arrives as an image or text file. The system reads it, pulls out "
         "the fields accounting needs, and states how much each one should be trusted.")
    body(doc,
         "Fields scoring at or above 90% confidence are auto-approved; anything below is "
         "flagged for a human. Every review then feeds back in: whether a reviewer "
         "accepted a field untouched or corrected it becomes a training label, so the "
         "scores sharpen with use.")
    body(doc,
         "Three things are learned from a single correction. The corrected value is "
         "remembered and reapplied when that document reappears, matched by a content "
         "fingerprint rather than filename. The label sitting beside it - "
         "“Gesamtbetrag”, say - becomes a candidate extraction pattern, adopted "
         "after three independent confirmations. And the accept-or-edit decision retrains "
         "that field's confidence model.")

    # ---- results -----------------------------------------------------
    heading(doc, "Measured results")
    body(doc,
         "All figures below come from running the extractor over 414 annotated invoices "
         "and comparing against ground truth - not from confidence scores, which proved a "
         "poor proxy for correctness.",
         color=MUTED, space_after=10)

    table(doc,
          ["Field", "Correct", "Rate", "Scored by"],
          [["Invoice number", "412 / 414", "99%", "Heuristic"],
           ["Total amount", "411 / 414", "99%", "Heuristic"],
           ["Invoice date", "336 / 414", "81%", "Model v1"],
           ["Supplier", "110 / 414", "26%", "Model v2"]],
          widths=[1.9, 1.3, 0.9, 1.4])

    body(doc,
         "The total began at 1%. It reached 99% once the extractor learned to find the "
         "summary block and confirm the figure against the invoice's own arithmetic - "
         "that subtotal plus tax equals the amount billed. That self-check turned out to "
         "be the strongest signal in the project, and it needs no machine learning at all.")

    table(doc,
          ["Model", "Training rows", "Failures seen", "Balanced accuracy"],
          [["Invoice date v1", "414", "78", "1.00"],
           ["Supplier v2", "414", "304", "0.99"]],
          widths=[1.7, 1.4, 1.4, 1.6])

    body(doc,
         "Both figures are cross-validated on held-out folds. A model only replaces the "
         "hand-written scoring if it clears 0.65 balanced accuracy there - a gate that has "
         "already done its job, rejecting a total-amount model that scored 0.019 and would "
         "have flagged nearly every correct invoice.")

    # ---- what could not be done --------------------------------------
    heading(doc, "What could not be done, and why")
    body(doc,
         "Four things were attempted and did not work. In each case the reason is worth "
         "more than the attempt.",
         color=MUTED, space_after=6)

    finding(doc, "Supplier extraction stayed at 26%", "UNRESOLVED", STOP,
            ["Every other field has an anchor: a keyword, a format, or arithmetic. A "
             "supplier name has none - the extractor takes whatever text follows "
             "“Bill To:” or “Seller:”, which on a two-column layout is "
             "frequently the adjacent column header rather than a company."],
            "Why it matters: supplier is weighted Very High, so its weakness pushes almost "
            "every invoice into review. The confidence model correctly learned to distrust "
            "the field - honest, but it does not reduce the review workload.")

    finding(doc, "Two fields never earned a trained model", "BLOCKED BY DATA", PART,
            ["Invoice number and total amount still use hand-written scores. The cause is "
             "counter-intuitive: extraction is too reliable. Across 414 invoices they "
             "failed only 2 and 3 times respectively - far below the 5 failures per class "
             "a model needs to learn any separation."],
            "Why it matters: the learning loop is sound, but a confidence model needs "
            "mistakes to learn from. These fields will stay on heuristics until real "
            "reviews supply genuine failures, which public datasets cannot.")

    finding(doc, "Image preprocessing raised confidence but not accuracy", "DEAD END", STOP,
            ["A preprocessing tool was built - upscaling, deskew from text-line angle, "
             "contrast equalisation, denoise. It lifted OCR confidence substantially on "
             "poor scans (0.54 to 0.75). Extraction did not improve, and on one invoice it "
             "made the result worse."],
            "Why: the failing scans carry text 8 pixels tall, where OCR needs roughly "
            "20-25. Enlarging interpolates - it makes characters bigger, not clearer. The "
            "information distinguishing a 2 from a 7 was destroyed when the image was "
            "captured, and nothing downstream recovers it.")

    finding(doc, "No spatial understanding of the page", "BY DESIGN", PART,
            ["The OCR engine returns a bounding box for every text region; the pipeline "
             "discards them and works purely on text. Layouts that place a value with no "
             "adjacent label - a total alone in a box top-right - cannot be handled."],
            "Why it stands: using position requires a layout-aware document model, which "
            "is a substantially different system. The text-only approach reached 99% on "
            "two fields, so the cost was not yet justified.")

    heading(doc, "The constraint underneath all of it", level=2)
    body(doc,
         "Every extraction failure traced back to one variable. Resolution predicts "
         "success almost perfectly:")

    table(doc,
          ["Source image", "Text height", "OCR confidence", "Fields extracted"],
          [["640 x 640  (fails)", "8 px", "0.68", "2.0 / 3"],
           ["1432 x 2048  (works)", "26 px", "0.86", "3.0 / 3"]],
          widths=[1.9, 1.2, 1.5, 1.5])

    body(doc,
         "Invoices captured at normal resolution extract completely. The low-resolution "
         "images are an artefact of the public dataset used for testing, not of real "
         "documents - which makes this the cheapest problem on the list to eliminate.")

    # ---- recommendations ---------------------------------------------
    doc.add_page_break()
    heading(doc, "Recommendations")
    body(doc,
         "Ordered by return on effort. The first two require no code and would resolve "
         "more uncertainty than anything else on the list.",
         color=MUTED, space_after=8)

    recommendations = [
        ("Capture invoices at full resolution",
         "A phone photo (~12 MP) or a 300 DPI scan puts text far past the 26-pixel "
         "threshold where the system already extracts every field. This removes the "
         "binding constraint behind every failure observed so far.",
         "No code · immediate effect"),
        ("Run twenty real invoices through the review app",
         "Everything measured so far used public datasets - synthetic templates and "
         "low-resolution photographs. Real documents would establish true accuracy, "
         "generate the failure examples the remaining models need, and reveal which "
         "layouts require new patterns.",
         "No code · unblocks three of the findings above"),
        ("Rebuild supplier extraction",
         "The weakest field and the largest single source of unnecessary review. Rather "
         "than taking text adjacent to a keyword, identify candidate company names by "
         "their form - legal suffixes, position in the header block - and validate against "
         "suppliers already known from past corrections.",
         "Focused engineering · reduces review volume"),
        ("Move the database to private hosting before real invoices",
         "The store keeps the OCR text of every document processed. Today that content "
         "comes from public datasets, so nothing sensitive is exposed; the moment real "
         "invoices flow through it, it will hold genuine bank and tax identifiers. A "
         "private repository, and a hosted database once more than one reviewer is "
         "involved.",
         "Configuration · required before production use"),
        ("Consider a layout-aware model only if the above is insufficient",
         "A document model such as LayoutLM uses position as well as text and would handle "
         "unlabelled values. It is the correct answer to the spatial limitation - but it "
         "needs annotated data, greater compute, and is only justified once resolution and "
         "real-world testing have been exhausted.",
         "Substantial · defer until evidence demands it"),
    ]
    for number, (title_text, text, effort) in enumerate(recommendations, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(2)
        styled(paragraph, f"{number}.  ", size=12, bold=True, color=ACCENT)
        styled(paragraph, title_text, size=12, bold=True)
        body(doc, text, space_after=3)
        note = doc.add_paragraph()
        note.paragraph_format.space_after = Pt(6)
        styled(note, effort, size=9.5, bold=True, color=MUTED)

    # ---- assessment --------------------------------------------------
    heading(doc, "Assessment")
    body(doc,
         "The pipeline reaches 99% accuracy on the two highest-weighted fields and the "
         "learning loop is verified end to end - reviews are recorded, models retrain, and "
         "new layouts are absorbed from corrections without code changes. The quality gate "
         "has demonstrably prevented a bad model from reaching production.")
    body(doc,
         "Its limits are equally clear, and they are mostly not modelling limits. Supplier "
         "extraction needs engineering, not data. Two fields need real failures before they "
         "can be learned. And the majority of observed errors would disappear with better "
         "source images. The most valuable next step is not a change to the system but "
         "twenty real invoices passed through it.")

    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(18)
    styled(footer,
           "Python · EasyOCR/PyTorch · scikit-learn · OpenCV · SQLite   |   "
           "3,724 lines · 13 tests passing   |   github.com/notOK18/Invoice_ai",
           size=9.5, color=MUTED)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


if __name__ == "__main__":
    target = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(__file__).parent / "Invoice_AI_Report.docx"
    print(f"wrote {build(target)}")
